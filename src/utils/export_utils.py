from docx import Document
from docx.shared import Pt
from xhtml2pdf import pisa


def export_to_docx(resume_text: str, cover_letter_text: str, output_path: str = "application.docx"):
    """Export resume + cover letter to DOCX format."""
    doc = Document()

    # Title
    title = doc.add_paragraph("Smart Resume & Cover Letter")
    title.runs[0].font.size = Pt(16)
    doc.add_paragraph()

    # Resume section
    doc.add_heading("Tailored Resume Summary", level=1)
    for line in resume_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    doc.add_page_break()

    # Cover letter section
    doc.add_heading("Cover Letter", level=1)
    for line in cover_letter_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    doc.save(output_path)
    return output_path


def export_to_pdf(resume_text: str, cover_letter_text: str, output_path: str = "application.pdf"):
    """Generate PDF using pure-Python xhtml2pdf (no external binary)."""

    # Handle newline conversion BEFORE building f-string
    resume_html = resume_text.replace("\n", "<br>")
    cover_letter_html = cover_letter_text.replace("\n", "<br>")

    html = f"""
    <html>
      <head>
        <meta charset="UTF-8">
        <style>
          body {{ font-family: Arial, sans-serif; margin: 40px; }}
          h2 {{ color: #2E3A59; }}
          hr {{ border: 0; border-top: 1px solid #ccc; margin: 20px 0; }}
          p {{ line-height: 1.5; }}
        </style>
      </head>
      <body>
        <h2>Tailored Resume Summary</h2>
        <p>{resume_html}</p>
        <hr>
        <h2>Cover Letter</h2>
        <p>{cover_letter_html}</p>
      </body>
    </html>
    """

    with open(output_path, "wb") as pdf_file:
        pisa_status = pisa.CreatePDF(html, dest=pdf_file)

    if pisa_status.err:
        raise Exception("Error generating PDF. Check HTML content.")
    return output_path